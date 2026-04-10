import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(r"D:\master_thesis")
TEX_PATH = ROOT / "thesis_proposal.tex"
BIB_PATH = ROOT / "references.bib"
OUT_PATH = ROOT / "thesis_proposal.docx"


def extract_braced_command(text: str, cmd: str) -> str:
    match = re.search(r"\\" + re.escape(cmd) + r"(?:\[[^\]]*\])?\{", text)
    if not match:
        return ""
    start = match.end() - 1
    depth = 0
    content_start = start
    for i in range(start, len(text)):
        ch = text[i]
        if ch == "{":
            depth += 1
            if depth == 1:
                content_start = i + 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return text[content_start:i]
    return ""


def clean_inline(s: str, cite_map: dict[str, int]) -> str:
    s = s.replace("~", " ")
    s = re.sub(
        r"\\cite\{([^}]+)\}",
        lambda m: "".join(f" [{cite_map.get(k.strip(), '?')}]" for k in m.group(1).split(",")),
        s,
    )
    s = re.sub(r"\\textbf\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\\label\{[^}]*\}", "", s)
    s = re.sub(r"\\[a-zA-Z]+\{([^}]*)\}", r"\1", s)
    s = s.replace("{", "").replace("}", "")
    s = s.replace("``", '"').replace("''", '"')
    s = s.replace("\u2019", "'")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def parse_author(raw: str, cite_map: dict[str, int]) -> str:
    first = re.search(r"\\fnm\{([^}]*)\}", raw)
    last = re.search(r"\\sur\{([^}]*)\}", raw)
    if first and last:
        return f"{first.group(1)} {last.group(1)}"
    return clean_inline(raw, cite_map)


def parse_affil(raw: str) -> str:
    parts = []
    for cmd in ["orgdiv", "orgname", "city", "country"]:
        parts.extend(re.findall(r"\\" + cmd + r"\{([^}]*)\}", raw))
    return ", ".join(part.strip() for part in parts if part.strip())


def parse_bib_entries(bib_text: str) -> list[tuple[str, str, str, str, str]]:
    entries = []
    chunks = re.split(r"(?=@\w+\{)", bib_text)
    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk:
            continue
        key_match = re.match(r"@\w+\{([^,]+),", chunk)
        key = key_match.group(1) if key_match else ""

        def field(name: str) -> str:
            match = re.search(r"\b" + re.escape(name) + r"\s*=\s*\{([^}]*)\}", chunk, flags=re.S)
            return match.group(1).strip() if match else ""

        author = field("author").replace(" and others", " et al.").replace(" and ", ", ")
        title = field("title").replace("{", "").replace("}", "")
        booktitle = field("booktitle")
        year = field("year")
        entries.append((key, author, title, booktitle, year))
    return entries


def flush_buffer(doc: Document, buffer: list[str], cite_map: dict[str, int]) -> None:
    text = clean_inline(" ".join(buffer), cite_map)
    if text:
        doc.add_paragraph(text)
    buffer.clear()


def main() -> None:
    tex = TEX_PATH.read_text(encoding="utf-8")
    bib = BIB_PATH.read_text(encoding="utf-8")

    keys = re.findall(r"@\w+\{([^,]+),", bib)
    cite_map = {key: i + 1 for i, key in enumerate(keys)}

    raw_title = extract_braced_command(tex, "title")
    raw_author = extract_braced_command(tex, "author")
    raw_affil = extract_braced_command(tex, "affil")
    raw_abstract = extract_braced_command(tex, "abstract")

    body_match = re.search(r"\\maketitle\s*(.*?)\\backmatter", tex, flags=re.S)
    body = body_match.group(1) if body_match else ""
    lines = [line.rstrip() for line in body.splitlines()]

    entries = parse_bib_entries(bib)

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    for style_name in ["Title", "Heading 1", "Heading 2"]:
        if style_name in styles:
            styles[style_name].font.name = "Times New Roman"

    title_para = doc.add_paragraph()
    title_para.style = "Title"
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.add_run(clean_inline(raw_title, cite_map))

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_para.add_run(parse_author(raw_author, cite_map)).bold = True

    affil_para = doc.add_paragraph()
    affil_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affil_para.add_run(parse_affil(raw_affil))

    doc.add_paragraph("")
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(clean_inline(raw_abstract, cite_map))

    buffer: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped or stripped.startswith("%%"):
            flush_buffer(doc, buffer, cite_map)
            continue
        if stripped.startswith(r"\section{"):
            flush_buffer(doc, buffer, cite_map)
            title = re.search(r"\\section\{([^}]*)\}", stripped).group(1)
            doc.add_heading(clean_inline(title, cite_map), level=1)
            continue
        if stripped.startswith(r"\subsection{"):
            flush_buffer(doc, buffer, cite_map)
            title = re.search(r"\\subsection\{([^}]*)\}", stripped).group(1)
            doc.add_heading(clean_inline(title, cite_map), level=2)
            continue
        if stripped.startswith(r"\begin{itemize}") or stripped.startswith(r"\end{itemize}"):
            flush_buffer(doc, buffer, cite_map)
            continue
        if stripped.startswith(r"\item"):
            flush_buffer(doc, buffer, cite_map)
            item_text = stripped[len(r"\item"):].strip()
            bullet = doc.add_paragraph(style="List Bullet")
            bullet.add_run(clean_inline(item_text, cite_map))
            continue
        if stripped.startswith("\\"):
            continue
        buffer.append(stripped)
    flush_buffer(doc, buffer, cite_map)

    doc.add_heading("References", level=1)
    for i, (_key, author, title, venue, year) in enumerate(entries, 1):
        ref = f"[{i}] {clean_inline(author, cite_map)}. {clean_inline(title, cite_map)}. {clean_inline(venue, cite_map)} ({clean_inline(year, cite_map)})."
        doc.add_paragraph(ref)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
